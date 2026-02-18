import time
import hashlib
from dataclasses import dataclass
from typing import List, Tuple, Iterable, Optional, Union

import requests
from docx import Document
from tqdm import tqdm

# =========================
# KONFIG
# =========================
LIBRETRANSLATE_URL = "http://localhost:5009/translate"
SOURCE_LANG = "id"
TARGET_LANG = "en"

# Batasi ukuran batch supaya stabil (Anda bisa naikkan jika mesin kuat)
MAX_TOTAL_CHARS_PER_BATCH = 20000   # total char gabungan dalam 1 request
MAX_ITEMS_PER_BATCH = 50           # jumlah paragraf per request

RETRY_LIMIT = 4
REQUEST_TIMEOUT = 180


@dataclass
class ParaRef:
    """Referensi paragraf yang akan ditulis balik."""
    paragraph: any
    original_text: str


def _is_blank(s: str) -> bool:
    return (s is None) or (s.strip() == "")


def _hash_text(s: str) -> str:
    return hashlib.sha1(s.encode("utf-8")).hexdigest()


def iter_all_paragraphs(doc: Document) -> Iterable[ParaRef]:
    """Ambil semua paragraf: body + tabel (cell paragraphs)."""
    for p in doc.paragraphs:
        yield ParaRef(p, p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield ParaRef(p, p.text)


def batcher(items: List[ParaRef],
            max_total_chars: int,
            max_items: int) -> Iterable[List[ParaRef]]:
    """Kelompokkan paragraf menjadi batch berdasarkan total char dan jumlah item."""
    batch = []
    total = 0

    for it in items:
        txt = it.original_text or ""
        n = len(txt)

        if _is_blank(txt):
            continue

        # Jika satu paragraf sangat panjang, pecah manual ke beberapa potong
        # (LibreTranslate bisa berat jika 1 item terlalu panjang)
        if n > max_total_chars:
            # flush batch yang sedang jalan
            if batch:
                yield batch
                batch, total = [], 0

            start = 0
            while start < n:
                chunk = txt[start:start + max_total_chars]
                # buat ParaRef "sementara" agar ditulis balik terakumulasi
                # kita tandai paragraph yang sama, nanti digabung saat writeback
                yield [ParaRef(it.paragraph, chunk)]
                start += max_total_chars
            continue

        # mulai batch baru jika melebihi batas
        if (len(batch) + 1 > max_items) or (total + n > max_total_chars):
            if batch:
                yield batch
            batch = [it]
            total = n
        else:
            batch.append(it)
            total += n

    if batch:
        yield batch


def libretranslate_translate_batch(texts: List[str],
                                  session: requests.Session) -> List[str]:
    """
    Translate batch via LibreTranslate.
    Mengirim JSON: {"q":[...], ...}
    Mengembalikan list hasil terjemahan sesuai urutan input.
    """
    payload = {
        "q": texts,
        "source": SOURCE_LANG,
        "target": TARGET_LANG,
        "format": "text",
    }

    last_err = None
    for attempt in range(1, RETRY_LIMIT + 1):
        try:
            r = session.post(
                LIBRETRANSLATE_URL,
                json=payload,
                timeout=REQUEST_TIMEOUT,
                headers={"Content-Type": "application/json"},
            )
            r.raise_for_status()
            data = r.json()

            # Beberapa instance mengembalikan "translatedText" sebagai list (batch)
            # atau string (single). Kita normalisasi ke list.
            tt = data.get("translatedText")
            if isinstance(tt, list):
                return [str(x) for x in tt]
            if isinstance(tt, str) and len(texts) == 1:
                return [tt]

            # fallback: kalau format tidak sesuai
            raise ValueError(f"Unexpected response shape: {data}")

        except Exception as e:
            last_err = e
            sleep_s = 1.5 * attempt
            print(f"[WARN] Batch gagal (attempt {attempt}/{RETRY_LIMIT}): {e} | sleep {sleep_s:.1f}s")
            time.sleep(sleep_s)

    # kalau gagal total, kembalikan teks asli
    print(f"[ERROR] Batch gagal total, pakai teks asli. Error terakhir: {last_err}")
    return texts


def clear_and_set_paragraph_text(paragraph, new_text: str):
    """
    Menulis balik teks ke paragraph dengan cara paling aman.
    Ini akan menghapus run-run lama (format inline hilang),
    tapi style paragraf (Heading, Normal, dll) tetap.
    """
    paragraph.text = new_text


def translate_docx(input_path: str, output_path: str):
    doc = Document(input_path)

    # Kumpulkan semua paragraf yang punya isi
    all_paras = list(iter_all_paragraphs(doc))
    translatable = [p for p in all_paras if not _is_blank(p.original_text)]

    with requests.Session() as session:
        pbar = tqdm(total=len(translatable), desc="Translating paragraphs")

        # Karena ada kasus paragraf super-panjang yang di-split jadi beberapa batch kecil,
        # kita perlu akumulasi hasil per paragraph (berdasarkan id object paragraph).
        accum = {}  # key: id(paragraph) -> list of translated chunks in order

        for batch in batcher(translatable, MAX_TOTAL_CHARS_PER_BATCH, MAX_ITEMS_PER_BATCH):
            batch_texts = [b.original_text for b in batch]
            translated_list = libretranslate_translate_batch(batch_texts, session)

            # Writeback / accumulate
            for bref, ttxt in zip(batch, translated_list):
                key = id(bref.paragraph)
                accum.setdefault(key, []).append(ttxt)

            pbar.update(len(batch))

        pbar.close()

        # Tulis balik hasil (gabungkan chunk jika paragraf pernah dipecah)
        for pref in translatable:
            key = id(pref.paragraph)
            if key in accum:
                final_text = "".join(accum[key])
                clear_and_set_paragraph_text(pref.paragraph, final_text)

    doc.save(output_path)
    print(f"Selesai. Output: {output_path}")


if __name__ == "__main__":
    # Ganti sesuai nama file Anda
    translate_docx("input.docx", "output_en.docx")
